from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Create new Word document
doc = Document()

# Add a cover page (blank)
doc.add_section()

# Add table of contents
doc.add_paragraph('Table of Contents')

# Add  sections
sections = ['PREFACE', 'EXECUTIVE SUMMARY', 'INTRODUCTION', 'LITERATURE REVIEW',
            'RESEARCH QUESTIONS', 'METHODOLOGY', 'DATA COLLECTION, ANALYSIS AND FINDINGS',
            'CONCLUSIONS & RECOMMENDATIONS', 'APPENDIX']

subsections = ['Research Methods', 'Desk Research', 'Field Research', 'Analysis of findings',
               'Works Cited', 'Time Frame / Schedule', 'Risks and Limitations']

#maybe some stylistic stuff
for section in sections:
    doc.add_paragraph(section, style='Heading 1')

for subsection in subsections:
    doc.add_paragraph(subsection, style='Heading 2')

# Save the document 
doc.save(r'Path/to/document/loc/here')